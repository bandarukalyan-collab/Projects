from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import sys

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def parse_inline_formatting(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    parse_inline_formatting(p, text)
    return p

def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    parse_inline_formatting(p, text)
    return p

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

md_file = sys.argv[1] if len(sys.argv) > 1 else 'RTR_GL_INT_009_Summary.md'
docx_file = sys.argv[2] if len(sys.argv) > 2 else 'RTR_GL_INT_009_Summary.docx'
md_dir = os.path.dirname(os.path.abspath(md_file))

with open(md_file, 'r', encoding='utf-8') as f:
    lines = f.read().split('\n')

in_code_block = False
code_lines = []
code_lang = ''
in_table = False
table_lines = []

for line in lines:
    stripped = line.rstrip('\n')
    
    if stripped.startswith('```'):
        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            code_lines = []
            in_code_block = False
        else:
            in_code_block = True
            code_lang = stripped[3:].strip()
        continue
    
    if in_code_block:
        if code_lang == 'mermaid' and stripped.startswith('flowchart'):
            code_lines.append('(Mermaid diagram - see markdown source)')
        else:
            code_lines.append(stripped)
        continue
    
    if stripped.startswith('!['):
        m = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if m:
            img_path = os.path.join(md_dir, m.group(2))
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(6.0))
        continue
    
    if stripped.startswith('|'):
        table_lines.append(stripped)
        in_table = True
        continue
    elif in_table:
        if table_lines:
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')]
                cells = [c for c in cells if c]
                if cells:
                    rows.append(cells)
            
            if rows:
                data_rows = [r for r in rows if not all(re.match(r'^[-:]+$', c) for c in r)]
                if data_rows:
                    max_cols = max(len(r) for r in data_rows)
                    table = doc.add_table(rows=1, cols=max_cols)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i in range(max_cols):
                        hdr_cells[i].text = data_rows[0][i] if i < len(data_rows[0]) else ''
                        set_cell_shading(hdr_cells[i], 'D9E2F3')
                    for row_data in data_rows[1:]:
                        row_cells = table.add_row().cells
                        for i in range(max_cols):
                            row_cells[i].text = row_data[i] if i < len(row_data) else ''
                    doc.add_paragraph()
            table_lines = []
        in_table = False
        continue
    
    if stripped.startswith('# '):
        h = doc.add_heading(stripped[2:], level=1)
        h.paragraph_format.keep_with_next = True
    elif stripped.startswith('## '):
        h = doc.add_heading(stripped[3:], level=2)
        h.paragraph_format.keep_with_next = True
    elif stripped.startswith('### '):
        h = doc.add_heading(stripped[4:], level=3)
        h.paragraph_format.keep_with_next = True
    elif stripped.startswith('#### '):
        h = doc.add_heading(stripped[5:], level=4)
        h.paragraph_format.keep_with_next = True
    elif re.match(r'^(\s*)[-*]\s+', stripped):
        match = re.match(r'^(\s*)[-*]\s+(.*)$', stripped)
        level = len(match.group(1)) // 2
        text = match.group(2)
        add_bullet(doc, text, level)
    elif re.match(r'^(\s*)(\d+)\.\s+(.*)$', stripped):
        match = re.match(r'^(\s*)(\d+)\.\s+(.*)$', stripped)
        level = len(match.group(1)) // 2
        text = match.group(3)
        add_numbered(doc, text, level)
    elif stripped.strip() == '':
        continue
    else:
        p = doc.add_paragraph()
        parse_inline_formatting(p, stripped)

doc.save(docx_file)
print('DOCX created:', docx_file)
